import asyncio
import hashlib
import json
import logging
import os
import re
import subprocess
import sys
import tempfile
import time
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qs, urlencode, urlparse
from urllib.request import Request, urlopen

import imageio_ffmpeg
import deno
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from dotenv import load_dotenv
from google import genai
from telegram import Update
from telegram.ext import Application, ContextTypes, MessageHandler, filters

load_dotenv()

TG_TOKEN = os.getenv("TG_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.7-flash")
GEMINI_JOB_TIMEOUT = int(os.getenv("GEMINI_JOB_TIMEOUT", "600"))
GEMINI_POLL_INTERVAL = int(os.getenv("GEMINI_POLL_INTERVAL", "5"))
SUPADATA_API_KEY = os.getenv("SUPADATA_API_KEY")
YOUTUBE_COOKIES_FILE = os.getenv(
    "YOUTUBE_COOKIES_FILE", "/etc/secrets/youtube_cookies.txt"
)
TELEGRAM_PROXY_URL = os.getenv("TELEGRAM_PROXY_URL")
PROCESSING_LOCK = asyncio.Lock()

logging.basicConfig(
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
    level=os.getenv("LOG_LEVEL", "INFO"),
)
LOGGER = logging.getLogger(__name__)
# HTTPX includes the complete Telegram Bot API URL in INFO messages. That URL
# contains the bot token, so third-party HTTP logs must never be emitted at INFO.
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("httpcore").setLevel(logging.WARNING)

SUMMARY_PROMPT = (Path(__file__).parent / "prompts" / "summary_ru.md").read_text(
    encoding="utf-8"
)


def normalize_youtube_url(raw_url: str) -> tuple[str, str]:
    """Return a canonical YouTube URL and its 11-character video id."""
    candidate = raw_url.strip()
    if not candidate.startswith(("http://", "https://")):
        candidate = f"https://{candidate}"

    parsed = urlparse(candidate)
    host = (parsed.hostname or "").lower()
    video_id = ""

    if host in {"youtu.be", "www.youtu.be"}:
        video_id = parsed.path.strip("/").split("/")[0]
    elif host == "youtube.com" or host.endswith(".youtube.com"):
        if parsed.path == "/watch":
            video_id = parse_qs(parsed.query).get("v", [""])[0]
        else:
            parts = [part for part in parsed.path.split("/") if part]
            if len(parts) >= 2 and parts[0] in {"shorts", "live", "embed"}:
                video_id = parts[1]
    else:
        raise ValueError("Пришли ссылку с youtube.com или youtu.be")

    if not re.fullmatch(r"[A-Za-z0-9_-]{11}", video_id):
        raise ValueError("Не удалось определить ID видео в ссылке")

    return f"https://www.youtube.com/watch?v={video_id}", video_id


def create_gemini_client():
    if not GEMINI_API_KEY:
        raise RuntimeError("На сервере не задан GEMINI_API_KEY")
    return genai.Client(api_key=GEMINI_API_KEY)


def fetch_supadata_transcript(url: str) -> str:
    """Fetch an existing Russian/available YouTube transcript without AI ASR."""
    if not SUPADATA_API_KEY:
        raise RuntimeError("На сервере не задан SUPADATA_API_KEY")

    query = urlencode(
        {
            "url": url,
            "lang": "ru",
            "text": "true",
            "mode": "native",
        }
    )
    request = Request(
        f"https://api.supadata.ai/v1/transcript?{query}",
        headers={"x-api-key": SUPADATA_API_KEY, "Accept": "application/json"},
    )
    LOGGER.info("Requesting native transcript from Supadata")
    try:
        with urlopen(request, timeout=60) as response:
            status_code = response.status
            payload = json.loads(response.read().decode("utf-8"))
    except HTTPError as error:
        details = error.read().decode("utf-8", errors="replace")[-800:]
        raise RuntimeError(f"Supadata вернул HTTP {error.code}: {details}") from error
    except (URLError, TimeoutError) as error:
        raise RuntimeError(f"Supadata временно недоступен: {error}") from error
    except (UnicodeDecodeError, json.JSONDecodeError) as error:
        raise RuntimeError("Supadata вернул некорректный ответ") from error

    content = payload.get("content")
    if isinstance(content, list):
        content = " ".join(
            str(item.get("text", "")) if isinstance(item, dict) else str(item)
            for item in content
        )
    transcript = str(content or "").strip()
    if status_code == 206 or not transcript:
        details = payload.get("message") or payload.get("error") or "субтитры отсутствуют"
        raise RuntimeError(f"Supadata не нашёл готовую расшифровку: {details}")

    LOGGER.info(
        "Supadata transcript received: language=%s, characters=%s",
        payload.get("lang", "unknown"),
        len(transcript),
    )
    return transcript


def request_summary(client, media_input: dict) -> str:
    """Run a long Gemini analysis in the background and poll its state."""
    interaction = None
    started_at = time.monotonic()
    deadline = started_at + GEMINI_JOB_TIMEOUT
    try:
        LOGGER.info("Creating background Gemini interaction")
        interaction = client.interactions.create(
            model=GEMINI_MODEL,
            input=[media_input, {"type": "text", "text": SUMMARY_PROMPT}],
            background=True,
            timeout=60,
        )
        interaction_id = interaction.id
        if not interaction_id:
            raise RuntimeError("Gemini не вернул ID фоновой задачи")

        last_status = None
        consecutive_poll_errors = 0
        while True:
            raw_status = getattr(interaction, "status", "")
            status_value = getattr(raw_status, "value", raw_status)
            status_name = str(status_value or "unknown").lower().split(".")[-1]
            if status_name != last_status:
                LOGGER.info(
                    "Gemini interaction %s status: %s", interaction_id, status_name
                )
                last_status = status_name

            if status_name == "completed":
                summary = (getattr(interaction, "output_text", "") or "").strip()
                if not summary:
                    raise RuntimeError("Gemini завершил задачу, но вернул пустой ответ")
                LOGGER.info(
                    "Gemini interaction %s completed in %s seconds",
                    interaction_id,
                    round(time.monotonic() - started_at),
                )
                return summary

            if status_name in {"failed", "cancelled", "canceled", "expired"}:
                details = getattr(interaction, "error", None) or status_name
                raise RuntimeError(f"задача Gemini завершилась со статусом {details}")

            if time.monotonic() >= deadline:
                try:
                    client.interactions.cancel(interaction_id, timeout=30)
                except Exception as cancel_error:
                    LOGGER.warning(
                        "Could not cancel Gemini interaction %s: %s",
                        interaction_id,
                        cancel_error,
                    )
                raise TimeoutError(
                    f"Gemini не завершил анализ за {GEMINI_JOB_TIMEOUT // 60} минут"
                )

            time.sleep(GEMINI_POLL_INTERVAL)
            try:
                interaction = client.interactions.get(interaction_id, timeout=30)
                consecutive_poll_errors = 0
            except Exception as poll_error:
                consecutive_poll_errors += 1
                LOGGER.warning(
                    "Gemini interaction %s poll failed (%s/5): %s",
                    interaction_id,
                    consecutive_poll_errors,
                    poll_error,
                )
                poll_error_text = str(poll_error).lower()
                if "error code: 400" in poll_error_text or "error code: 403" in poll_error_text:
                    raise RuntimeError(
                        "Gemini не смог получить доступ к этому видео "
                        f"({poll_error})"
                    ) from poll_error
                if consecutive_poll_errors >= 5:
                    raise RuntimeError(
                        "Не удалось получить состояние задачи Gemini после 5 попыток"
                    ) from poll_error
    except Exception:
        if interaction is None:
            LOGGER.exception("Could not create background Gemini interaction")
        raise


def summarize_youtube_url(client, url: str) -> str:
    return request_summary(
        client,
        {"type": "video", "uri": url, "mime_type": "video/mp4"},
    )


def summarize_transcript(client, transcript: str) -> str:
    """Summarize transcript text with a regular, bounded Gemini request."""
    LOGGER.info("Sending Supadata transcript to Gemini")
    response = client.interactions.create(
        model=GEMINI_MODEL,
        input=[
            {
                "type": "text",
                "text": (
                    "Ниже приведена расшифровка русскоязычного YouTube-видео. "
                    "Используй её как единственный источник содержания.\n\n"
                    f"{transcript}"
                ),
            },
            {"type": "text", "text": SUMMARY_PROMPT},
        ],
        timeout=300,
    )
    summary = (response.output_text or "").strip()
    if not summary:
        raise RuntimeError("Gemini вернул пустой ответ на расшифровку")
    return summary


def download_audio(url: str, destination: Path) -> Path:
    """Download and convert one video's audio to an API-supported MP3 file."""
    output_template = str(destination / "audio.%(ext)s")
    ffmpeg_dir = str(Path(imageio_ffmpeg.get_ffmpeg_exe()).parent)
    deno_path = deno.find_deno_bin()
    command = [
        sys.executable,
        "-m",
        "yt_dlp",
        "--no-playlist",
        "--retries",
        "3",
        "--fragment-retries",
        "3",
        "--socket-timeout",
        "30",
        "--js-runtimes",
        f"deno:{deno_path}",
        "--extract-audio",
        "--audio-format",
        "mp3",
        "--audio-quality",
        "5",
        "--ffmpeg-location",
        ffmpeg_dir,
        "--output",
        output_template,
    ]
    cookies_path = Path(YOUTUBE_COOKIES_FILE)
    if cookies_path.is_file():
        LOGGER.info("Using YouTube cookies from configured secret file")
        command.extend(["--cookies", str(cookies_path)])
    command.append(url)
    result = subprocess.run(
        command,
        capture_output=True,
        text=True,
        timeout=900,
        check=False,
    )
    audio_path = destination / "audio.mp3"
    if result.returncode != 0 or not audio_path.is_file():
        details = (result.stderr or result.stdout)[-800:].strip()
        raise RuntimeError(f"Не удалось скачать аудио с YouTube: {details}")
    return audio_path


def summarize_audio_fallback(client, url: str, destination: Path) -> str:
    audio_path = download_audio(url, destination)
    uploaded_file = None
    try:
        uploaded_file = client.files.upload(file=str(audio_path))
        return request_summary(
            client,
            {
                "type": "audio",
                "uri": uploaded_file.uri,
                "mime_type": uploaded_file.mime_type or "audio/mp3",
            },
        )
    finally:
        if uploaded_file is not None:
            try:
                client.files.delete(name=uploaded_file.name)
            except Exception as error:
                LOGGER.warning("Could not delete temporary Gemini file: %s", error)


def generate_summary(url: str, destination: Path) -> tuple[str, bool]:
    """Prefer a native transcript, then direct video, then downloaded audio."""
    client = create_gemini_client()
    try:
        if SUPADATA_API_KEY:
            try:
                transcript = fetch_supadata_transcript(url)
                return summarize_transcript(client, transcript), False
            except Exception as transcript_error:
                LOGGER.warning(
                    "Supadata transcript path failed, using direct video: %s",
                    transcript_error,
                )
        else:
            LOGGER.warning("SUPADATA_API_KEY is not configured; skipping transcript")
        try:
            LOGGER.info("Starting direct YouTube analysis for %s", url)
            return summarize_youtube_url(client, url), False
        except Exception as direct_error:
            LOGGER.warning("Direct YouTube analysis failed, using audio: %s", direct_error)
            try:
                return summarize_audio_fallback(client, url, destination), True
            except Exception as audio_error:
                raise RuntimeError(
                    "Gemini не получил доступ к видео, а резервное скачивание "
                    f"аудио тоже не удалось. Gemini: {direct_error}. "
                    f"YouTube: {audio_error}"
                ) from audio_error
    finally:
        try:
            client.close()
        except Exception as close_error:
            LOGGER.warning("Could not close Gemini client: %s", close_error)


def add_inline_markdown(paragraph, text: str):
    """Render the small bold subset requested from Gemini without raw ** marks."""
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def add_markdown_to_document(document: Document, markdown: str):
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            document.add_heading(heading.group(2).strip(" *"), level=len(heading.group(1)))
            continue

        bullet = re.match(r"^[-*•]\s+(.+)$", line)
        paragraph = document.add_paragraph(style="List Bullet" if bullet else None)
        add_inline_markdown(paragraph, bullet.group(1) if bullet else line)


def configure_document(document: Document):
    """Use the compact A4 layout and restrained hierarchy of the reference."""
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, before, after in (
        ("Heading 1", 12, 12, 6),
        ("Heading 2", 11, 8, 3),
        ("Heading 3", 11, 6, 3),
    ):
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def create_document(summary: str, filename: Path):
    document = Document()
    configure_document(document)
    add_markdown_to_document(document, summary)
    document.save(filename)


async def handle_link(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.message is None or not update.message.text:
        return

    try:
        url, video_id = normalize_youtube_url(update.message.text)
    except ValueError as error:
        await update.message.reply_text(f"❌ {error}")
        return

    status = await update.message.reply_text("⏳ Видео принято, ожидаю обработку…")
    async with PROCESSING_LOCK:
        await status.edit_text("🔎 Получаю субтитры и создаю саммари…")
        try:
            with tempfile.TemporaryDirectory(prefix=f"summary-{video_id}-") as temp_dir:
                temp_path = Path(temp_dir)
                LOGGER.info("Accepted video %s for processing", video_id)
                started_at = time.monotonic()
                processing_task = asyncio.create_task(
                    asyncio.to_thread(generate_summary, url, temp_path)
                )
                while True:
                    done, _ = await asyncio.wait({processing_task}, timeout=60)
                    if processing_task in done:
                        summary, used_audio = await processing_task
                        break
                    elapsed_minutes = max(1, int((time.monotonic() - started_at) // 60))
                    LOGGER.info(
                        "Video %s is still processing (%s min)",
                        video_id,
                        elapsed_minutes,
                    )
                    try:
                        await status.edit_text(
                            "🤖 Видео обрабатывается…\n"
                            f"⏱ Прошло: {elapsed_minutes} мин."
                        )
                    except Exception as status_error:
                        LOGGER.warning("Could not update Telegram status: %s", status_error)
                if used_audio:
                    await status.edit_text("📄 Аудио распознано, создаю Word-файл…")
                else:
                    await status.edit_text("📄 Видео проанализировано, создаю Word-файл…")

                filename = temp_path / f"summary_{video_id}.docx"
                await asyncio.to_thread(create_document, summary, filename)
                with filename.open("rb") as document:
                    await update.message.reply_document(
                        document=document,
                        filename=filename.name,
                        caption="✅ Саммари готово!",
                    )
            await status.delete()
        except Exception as error:
            LOGGER.exception("Failed to process %s", url)
            message = str(error)
            if len(message) > 900:
                message = message[:900] + "…"
            await status.edit_text(f"❌ Не удалось обработать видео: {message}")


def validate_configuration():
    missing = [
        name
        for name, value in (("TG_TOKEN", TG_TOKEN), ("GEMINI_API_KEY", GEMINI_API_KEY))
        if not value
    ]
    if missing:
        raise RuntimeError(f"Не заданы обязательные переменные: {', '.join(missing)}")


def webhook_security_values(token: str) -> tuple[str, str]:
    """Derive a non-secret URL path and Telegram header secret from the token."""
    digest = hashlib.sha256(token.encode("utf-8")).hexdigest()
    return f"telegram/{digest[:24]}", digest[24:]


def main():
    validate_configuration()
    LOGGER.info("Starting bot with model %s", GEMINI_MODEL)
    builder = (
        Application.builder()
        .token(TG_TOKEN)
        .connect_timeout(30)
        .get_updates_connect_timeout(30)
    )
    if TELEGRAM_PROXY_URL:
        builder = builder.proxy(TELEGRAM_PROXY_URL).get_updates_proxy(
            TELEGRAM_PROXY_URL
        )
    application = builder.build()
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_link))
    external_url = os.getenv("RENDER_EXTERNAL_URL") or os.getenv("WEBHOOK_URL")
    if external_url:
        webhook_path, webhook_secret = webhook_security_values(TG_TOKEN)
        LOGGER.info("Starting Telegram webhook")
        application.run_webhook(
            listen="0.0.0.0",
            port=int(os.getenv("PORT", 10000)),
            url_path=webhook_path,
            webhook_url=f"{external_url.rstrip('/')}/{webhook_path}",
            secret_token=webhook_secret,
            bootstrap_retries=3,
        )
    else:
        LOGGER.info("Starting Telegram polling")
        application.run_polling(bootstrap_retries=3)


if __name__ == "__main__":
    main()
