import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import Mock, patch

import main


class NormalizeYoutubeUrlTests(unittest.TestCase):
    def test_supported_urls(self):
        expected = "https://www.youtube.com/watch?v=dQw4w9WgXcQ"
        for url in (
            "https://youtu.be/dQw4w9WgXcQ?t=10",
            "https://www.youtube.com/watch?v=dQw4w9WgXcQ&list=abc",
            "https://youtube.com/shorts/dQw4w9WgXcQ",
            "https://m.youtube.com/live/dQw4w9WgXcQ?feature=share",
        ):
            with self.subTest(url=url):
                self.assertEqual(
                    main.normalize_youtube_url(url),
                    (expected, "dQw4w9WgXcQ"),
                )

    def test_rejects_non_youtube_and_invalid_id(self):
        for url in ("https://example.com/video", "https://youtu.be/short"):
            with self.subTest(url=url), self.assertRaises(ValueError):
                main.normalize_youtube_url(url)

    def test_webhook_values_are_stable_and_do_not_expose_token(self):
        token = "123456:secret-token"
        path, secret = main.webhook_security_values(token)
        self.assertEqual((path, secret), main.webhook_security_values(token))
        self.assertNotIn(token, path)
        self.assertNotIn(token, secret)
        self.assertTrue(path.startswith("telegram/"))


class SummaryPipelineTests(unittest.TestCase):
    @patch("main.urlopen")
    def test_supadata_uses_native_russian_transcript(self, urlopen):
        response = Mock()
        response.status = 200
        response.read.return_value = json_bytes = (
            b'{"content":"Russian transcript","lang":"ru"}'
        )
        response.__enter__ = Mock(return_value=response)
        response.__exit__ = Mock(return_value=False)
        urlopen.return_value = response

        with patch.object(main, "SUPADATA_API_KEY", "supadata-key"):
            transcript = main.fetch_supadata_transcript("https://youtu.be/example")

        self.assertEqual(transcript, "Russian transcript")
        request = urlopen.call_args.args[0]
        self.assertIn("mode=native", request.full_url)
        self.assertIn("lang=ru", request.full_url)
        self.assertEqual(json_bytes, response.read.return_value)

    def test_supadata_is_preferred_before_direct_video(self):
        client = Mock()
        with patch.object(main, "SUPADATA_API_KEY", "supadata-key"), patch.object(
            main, "create_gemini_client", return_value=client
        ), patch.object(
            main, "fetch_supadata_transcript", return_value="расшифровка"
        ) as fetch, patch.object(
            main, "summarize_transcript", return_value="готово"
        ) as summarize, patch.object(main, "summarize_youtube_url") as direct:
            result = main.generate_summary("https://youtu.be/dQw4w9WgXcQ", Path("."))

        self.assertEqual(result, ("готово", False))
        fetch.assert_called_once()
        summarize.assert_called_once_with(client, "расшифровка")
        direct.assert_not_called()

    @patch("main.time.sleep")
    def test_transcript_summary_retries_temporary_gemini_quota(self, sleep):
        client = Mock()
        client.interactions.create.side_effect = [
            RuntimeError("429: Please retry in 24.1s"),
            SimpleNamespace(output_text="готовое саммари"),
        ]

        result = main.summarize_transcript(client, "расшифровка")

        self.assertEqual(result, "готовое саммари")
        self.assertEqual(client.interactions.create.call_count, 2)
        self.assertEqual(
            client.interactions.create.call_args.kwargs["model"],
            main.GEMINI_TEXT_MODEL,
        )
        sleep.assert_called_once_with(27.1)

    def test_gemini_text_failure_does_not_refetch_video(self):
        client = Mock()
        with patch.object(main, "SUPADATA_API_KEY", "supadata-key"), patch.object(
            main, "create_gemini_client", return_value=client
        ), patch.object(
            main, "fetch_supadata_transcript", return_value="расшифровка"
        ), patch.object(
            main, "summarize_transcript", side_effect=RuntimeError("429 quota")
        ), patch.object(main, "summarize_youtube_url") as direct:
            with self.assertRaisesRegex(RuntimeError, "Расшифровка получена"):
                main.generate_summary("https://youtu.be/dQw4w9WgXcQ", Path("."))

        direct.assert_not_called()

    @patch("main.time.sleep")
    def test_summary_uses_background_interaction(self, _sleep):
        client = Mock()
        client.interactions.create.return_value = SimpleNamespace(
            id="interaction-1", status="in_progress", output_text=""
        )
        client.interactions.get.return_value = SimpleNamespace(
            id="interaction-1", status="completed", output_text="готовое саммари"
        )

        result = main.request_summary(client, {"type": "video", "uri": "https://x"})

        self.assertEqual(result, "готовое саммари")
        self.assertTrue(client.interactions.create.call_args.kwargs["background"])
        client.interactions.get.assert_called_once_with("interaction-1", timeout=30)

    def test_summary_reports_failed_background_interaction(self):
        client = Mock()
        client.interactions.create.return_value = SimpleNamespace(
            id="interaction-1",
            status="failed",
            output_text="",
            error="RESOURCE_EXHAUSTED",
        )

        with self.assertRaisesRegex(RuntimeError, "RESOURCE_EXHAUSTED"):
            main.request_summary(client, {"type": "video", "uri": "https://x"})

    @patch("main.subprocess.run")
    @patch("main.deno.find_deno_bin", return_value="C:/deno/deno.exe")
    @patch("main.imageio_ffmpeg.get_ffmpeg_exe", return_value="C:/ffmpeg/ffmpeg.exe")
    def test_downloader_uses_current_python(self, _ffmpeg, _deno, run):
        run.return_value.returncode = 1
        run.return_value.stderr = "download failed"
        with tempfile.TemporaryDirectory() as directory, self.assertRaises(RuntimeError):
            main.download_audio("https://youtu.be/dQw4w9WgXcQ", Path(directory))
        command = run.call_args.args[0]
        self.assertEqual(command[:3], [main.sys.executable, "-m", "yt_dlp"])
        self.assertIn("deno:C:/deno/deno.exe", command)

    def test_direct_youtube_is_preferred(self):
        client = Mock()
        with patch.object(main, "create_gemini_client", return_value=client), patch.object(
            main, "summarize_youtube_url", return_value="готово"
        ) as direct, patch.object(main, "summarize_audio_fallback") as fallback:
            result = main.generate_summary("https://youtu.be/dQw4w9WgXcQ", Path("."))
        self.assertEqual(result, ("готово", False))
        direct.assert_called_once()
        fallback.assert_not_called()

    def test_audio_is_used_when_direct_call_fails(self):
        client = Mock()
        with patch.object(main, "create_gemini_client", return_value=client), patch.object(
            main, "summarize_youtube_url", side_effect=RuntimeError("direct failed")
        ), patch.object(main, "summarize_audio_fallback", return_value="резерв") as fallback:
            result = main.generate_summary("https://youtu.be/dQw4w9WgXcQ", Path("."))
        self.assertEqual(result, ("резерв", True))
        fallback.assert_called_once()

    def test_document_is_created(self):
        with tempfile.TemporaryDirectory() as directory:
            target = Path(directory) / "summary.docx"
            main.create_document(
                "# 📌 О чем выпуск\n## 1. Модель\n- **Вывод:** Первый тезис",
                target,
            )
            self.assertTrue(target.is_file())
            self.assertGreater(target.stat().st_size, 0)
            from docx import Document

            created = Document(target)
            full_text = "\n".join(paragraph.text for paragraph in created.paragraphs)
            self.assertNotIn("**", full_text)
            self.assertEqual(created.paragraphs[0].style.name, "Heading 1")
            self.assertEqual(created.paragraphs[1].style.name, "Heading 2")
            self.assertTrue(created.paragraphs[2].runs[0].bold)
            self.assertAlmostEqual(created.sections[0].page_width.cm, 21, places=1)
            self.assertAlmostEqual(created.sections[0].left_margin.cm, 3, places=1)

    def test_prompt_matches_reference_structure(self):
        self.assertIn("# 📌 О чем этот выпуск", main.SUMMARY_PROMPT)
        self.assertIn("# 🧠 Ключевые модели и техники", main.SUMMARY_PROMPT)
        self.assertIn("Не включай рекламные интеграции", main.SUMMARY_PROMPT)
        self.assertIn("Не создавай отдельные разделы с таймкодами", main.SUMMARY_PROMPT)


if __name__ == "__main__":
    unittest.main()
